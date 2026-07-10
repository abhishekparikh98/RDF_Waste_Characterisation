"""
Build the complete MSc Project Study Guide as a .docx file.

This script generates a comprehensive study guide based on the actual project
files in the repository. It does NOT invent any content - everything comes from
the project source code, reports, logs, and configuration files.
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

OUT_PATH = r"D:\University\Msc Project\MSC_Project_Complete_Study_Guide.docx"

doc = Document()

# Page setup
for section in doc.sections:
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# Default style
styles = doc.styles
normal = styles['Normal']
normal.font.name = 'Calibri'
normal.font.size = Pt(11)


def add_heading(text, level=1, color=None):
    h = doc.add_heading(text, level=level)
    if color:
        for run in h.runs:
            run.font.color.rgb = color
    return h


def add_para(text, bold=False, italic=False, size=11, color=None, align=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color
    if align:
        p.alignment = align
    return p


def add_qa(question, answer):
    """Add a Question + bold answer pair."""
    qp = doc.add_paragraph()
    qp.add_run("Q: ").bold = True
    qp.add_run(question).bold = True
    ap = doc.add_paragraph()
    ap.add_run("A: ").bold = True
    ap.add_run(answer)
    return ap


def add_code(text):
    """Add a monospace block (rendered as a styled paragraph)."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    # Add a light background by setting shading
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'F4F4F4')
    pPr.append(shd)
    return p


def add_bullet(text):
    p = doc.add_paragraph(text, style='List Bullet')
    return p


def add_numbered(text):
    p = doc.add_paragraph(text, style='List Number')
    return p


def page_break():
    doc.add_page_break()


# =====================================================================
# COVER PAGE
# =====================================================================
title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title_p.add_run("MSC PROJECT")
r.bold = True
r.font.size = Pt(28)
r.font.color.rgb = RGBColor(0x12, 0x35, 0x5b)

sub_p = doc.add_paragraph()
sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub_p.add_run("Complete Study Guide & Viva Preparation")
r.font.size = Pt(20)
r.italic = True
r.font.color.rgb = RGBColor(0x33, 0x5C, 0x81)

doc.add_paragraph()

proj_p = doc.add_paragraph()
proj_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = proj_p.add_run("Multi-Modal Waste Characterisation for\nHigh-Quality Refuse-Derived Fuel Production\nUsing Machine Learning")
r.font.size = Pt(16)
r.bold = True

doc.add_paragraph()
doc.add_paragraph()

author_p = doc.add_paragraph()
author_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = author_p.add_run("Author: Abhishek Parikh (abhishekparikh98)")
r.font.size = Pt(13)
r.bold = True

inst_p = doc.add_paragraph()
inst_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = inst_p.add_run("MSc Computing Dissertation")
r.font.size = Pt(12)

year_p = doc.add_paragraph()
year_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = year_p.add_run("Year: 2026")
r.font.size = Pt(12)

doc.add_paragraph()
note_p = doc.add_paragraph()
note_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = note_p.add_run("This study guide is generated from the actual project repository.\nEvery fact, metric, file path, and code reference can be verified in the source code.")
r.italic = True
r.font.size = Pt(10)
r.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

page_break()

# =====================================================================
# TABLE OF CONTENTS (manual)
# =====================================================================
add_heading("Table of Contents", level=1, color=RGBColor(0x12, 0x35, 0x5b))
toc_items = [
    "Chapter 1: Project Overview",
    "Chapter 2: Repository Walkthrough",
    "Chapter 3: Scripts (How to Run the Project)",
    "Chapter 4: Datasets (TrashNet, TACO, Synthetic RDF)",
    "Chapter 5: Machine Learning Concepts (From Zero to Advanced)",
    "Chapter 6: Algorithms (CNN, MobileNetV2, ResNet50, Random Forest)",
    "Chapter 7: End-to-End Pipeline",
    "Chapter 8: Flask Web Application",
    "Chapter 9: Results, Graphs & Confusion Matrices",
    "Chapter 10: Technology Stack",
    "Chapter 11: Project Architecture",
    "Chapter 12: 150+ Viva Questions with Model Answers",
    "Chapter 13: Supervisor Questions & Realistic Answers",
    "Chapter 14: Presentation Notes (Slide by Slide)",
    "Chapter 15: Code Walkthrough (Line by Line)",
    "Chapter 16: Final Revision Notes, Memory Tricks, Cheat Sheet",
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.left_indent = Cm(0.3)
    p.paragraph_format.space_after = Pt(2)

page_break()

# =====================================================================
# CHAPTER 1: PROJECT OVERVIEW
# =====================================================================
add_heading("Chapter 1: Project Overview", level=1, color=RGBColor(0x12, 0x35, 0x5b))

add_heading("1.1 Project Title (Verbatim)", level=2)
add_para("Multi-Modal Waste Characterisation for High-Quality Refuse-Derived Fuel Production Using Machine Learning.")
add_para("Source: README.md line 1, pyproject.toml line 8.", italic=True, size=10, color=RGBColor(0x80, 0x80, 0x80))

add_heading("1.2 What is RDF (Refuse-Derived Fuel)?", level=2)
add_para("RDF stands for Refuse-Derived Fuel. It is a fuel produced from municipal solid waste (MSW) after a series of treatment steps:")
add_bullet("Step 1 - Sorting: separating recyclable and non-recyclable waste.")
add_bullet("Step 2 - Shredding: reducing the waste to small, uniform particles.")
add_bullet("Step 3 - Drying: removing moisture (a key factor for calorific value).")
add_bullet("Step 4 - Pelletising: compressing into pellets or fluff that can be burned in cement kilns or dedicated power plants.")
add_para("RDF is important because it diverts waste from landfill, reduces methane emissions, and produces energy. The quality of RDF depends on its calorific value, moisture content, contamination, and combustibility - exactly the four numerical features your model predicts.")

add_heading("1.3 Why Waste Classification Matters for RDF", level=2)
add_para("In a real Materials Recovery Facility (MRF), the input to the RDF line is a mixed stream. If you accidentally put glass or metal into the RDF pellet, you contaminate the fuel, damage the burners, and lower the calorific value. So the FIRST decision point is: which incoming item goes to RDF and which goes elsewhere (recycling, landfill, composting)? That decision is exactly what your model automates.")

add_heading("1.4 Why Machine Learning", level=2)
add_para("Manual sorting is slow (typically 30-60 kg/hour per worker) and inconsistent (human fatigue, training variation). Computer vision can run at conveyor-belt speeds (>1 item/second) and is consistent 24/7. By combining vision with a tabular model that knows about RDF chemistry, you get a system that does not just say 'this is plastic' but also 'this is good RDF material'.")

add_heading("1.5 Project Objectives (Verbatim from README.md)", level=2)
add_numbered("Develop machine learning models capable of classifying waste materials from multiple modalities (images, sensor data).")
add_numbered("Combine TrashNet and TACO datasets for comprehensive waste material characterisation.")
add_numbered("Evaluate model performance on waste-to-RDF production pipeline optimisation.")
add_numbered("Create interpretable models suitable for industrial deployment.")

add_heading("1.6 Scope (What the Project Does and Does Not Do)", level=2)
add_para("In scope:")
add_bullet("Six-class waste image classification: cardboard, glass, metal, paper, plastic, trash.")
add_bullet("Three image models compared: from-scratch CNN, MobileNetV2 transfer learning, ResNet50 transfer learning.")
add_bullet("One tabular model: Random Forest for RDF suitability.")
add_bullet("Multi-modal pipeline that chains image classification to RDF prediction.")
add_bullet("A Flask web demo for live inference.")
add_bullet("A command-line inference tool.")
add_bullet("Comprehensive evaluation, confusion matrices, training curves, and feature importance plots.")
add_para("Out of scope (be honest about these in the viva):")
add_bullet("TACO dataset is bundled as a library, but the actual TACO images are NOT downloaded and NOT used in the training pipeline. This is documented in the dataset report and the portfolio.")
add_bullet("The RDF tabular dataset is SYNTHETIC, generated from domain rules (moisture, contamination, combustibility, calorific_value). It is not real plant data.")
add_bullet("No cross-validation for the image models (single 70/15/15 stratified split).")
add_bullet("No data augmentation (rotation, flip, zoom) is applied.")
add_bullet("No deployment configuration (Docker, Gunicorn, HTTPS) - the Flask app is a local demo only.")
add_bullet("No authentication, no batch upload, no REST API beyond the single / endpoint.")

add_heading("1.7 Expected Outcomes", level=2)
add_bullet("Trained Keras models saved as .h5 files in models/.")
add_bullet("A trained scikit-learn Pipeline (.joblib) for RDF suitability.")
add_bullet("Markdown evaluation reports in reports/.")
add_bullet("300 DPI PNG figures in results/ and reports/figures/.")
add_bullet("A working Flask web app and a CLI inference tool.")
add_bullet("A reproducible random seed of 42 across all experiments.")
add_bullet("Test accuracy of 0.8877 (ResNet50) and 0.9133 (RDF Random Forest) on held-out test sets.")

add_heading("1.8 Quick Project Identity Card (Memorise This)", level=2)
add_bullet("Python: 3.11.9 (declared 3.9+ in pyproject.toml).")
add_bullet("TensorFlow: 2.21.0.")
add_bullet("scikit-learn: >=1.0.")
add_bullet("Random seed: 42 throughout.")
add_bullet("Training dates: 2026-06-26 to 2026-07-02 (5 git commits).")
add_bullet("Author: Abhishek Parikh (git user.name = abhishekparikh98).")
add_bullet("License: MIT.")
add_bullet("Pipeline: Image -> CNN -> Class -> Material features -> Random Forest -> RDF Suitable / Not Suitable.")

page_break()

# =====================================================================
# CHAPTER 2: REPOSITORY WALKTHROUGH
# =====================================================================
add_heading("Chapter 2: Repository Walkthrough", level=1, color=RGBColor(0x12, 0x35, 0x5b))

add_para("This chapter walks you through every folder and every file in the project. The goal is that after reading it, you can navigate the repository in your head without looking at your screen.", italic=True)

add_heading("2.1 Top-Level Files", level=2)
add_bullet("README.md - Project landing page. Contains title, objectives, structure, install instructions, and how to run the Flask app.")
add_bullet("PROJECT_EXECUTION_GUIDE.txt - Plain-text operator's manual. 157 lines. Walks you through the seven steps to run the project end-to-end (install, preprocess, train CNN, train comparison, train RF, run inference, run Flask).")
add_bullet("PROJECT_ANALYSIS.md - 533-line pre-existing audit. Lists every folder, every file, the technology stack, and a list of 'what is missing' (e.g. TACO is bundled but not used; notebooks/ is empty).")
add_bullet("requirements.txt - Pinned ranges for all Python dependencies. Includes numpy, pandas, scikit-learn, Flask, TensorFlow, Keras, OpenCV, Pillow, matplotlib, seaborn, plotly, plus tooling (pytest, black, flake8, sphinx).")
add_bullet("pyproject.toml - Build configuration. Project name is 'msc-waste-characterisation', version '0.1.0', license MIT, requires Python >=3.9.")
add_bullet(".env.example - Template for environment variables (RAW_DATA_PATH, PROCESSED_DATA_PATH, MODEL_SAVE_PATH, WANDB_PROJECT, RANDOM_SEED=42, LOG_LEVEL=INFO).")
add_bullet(".gitignore - Excludes __pycache__/, .venv/, models/checkpoints/, results/experiments/, .wandb/, runs/, mlruns/, *.tmp.")
add_bullet("LICENSE - MIT license, copyright 2026.")
add_bullet("app.py - Flask web app. 113 lines. Single route '/' with GET (render empty form) and POST (accept uploaded image, run inference, render result).")
add_bullet("training.log, preprocessing.log, exploration.log, comparison.log, rdf_training.log - Five execution logs. Each one records the timestamp, what was done, and any errors.")

add_heading("2.2 The src/ Folder (Core Library)", level=2)
add_para("Even though the README and pyproject.toml suggest sub-packages (src/models/, src/preprocessing/, etc.), the actual implementation uses flat modules directly in src/. This is one of the small inconsistencies between docs and code. Be honest about it in the viva.", italic=True)

add_heading("2.2.1 src/__init__.py", level=3)
add_para("Empty package marker. Defines __version__ = '0.1.0' and __author__ placeholder.")
add_code('__version__ = "0.1.0"\n__author__ = "[Your Name]"\n__description__ = "MSc dissertation project on waste characterization using machine learning"')

add_heading("2.2.2 src/config.py - Central Configuration", level=3)
add_para("Defines four dataclasses that act as the project's single source of truth for hyperparameters:")
add_bullet("DataConfig - image height/width (224x224), batch size (32), num classes (6), class_names = ['cardboard','glass','metal','paper','plastic','trash'].")
add_bullet("ModelConfig - num_classes, input_shape, dropout_rate (0.5), conv filters per block (32, 64, 128), dense_units (256).")
add_bullet("TrainingConfig - learning_rate (0.001), epochs (30), early_stopping_patience (5), early_stopping_min_delta (0.001), random_seed (42), optimizer ('adam'), loss_fn ('categorical_crossentropy'), metrics (['accuracy']).")
add_bullet("ExperimentConfig - experiment_name, project_name, seed (42), model_save_path, results_dir, reports_dir, log_file.")
add_para("All scripts import these defaults via from src.config import DEFAULT_DATA_CONFIG, ... This means if you change a hyperparameter in one place, it propagates to every script.")

add_heading("2.2.3 src/models.py - Model Architectures", level=3)
add_para("Defines four model factories:")
add_bullet("BaselineCNN class - a Sequential Keras model. Three Conv2D blocks (32, 64, 128 filters, 3x3 kernel, ReLU, same padding) -> MaxPool 2x2 -> Dropout 0.5 each. Then Flatten, Dense(256, ReLU), Dropout(0.5), Dense(6, Softmax). Total params: 25,785,158 trainable, 77,355,476 total (the difference is Adam optimizer state).")
add_bullet("build_baseline_cnn() - factory function that returns a compiled Keras model with the architecture above.")
add_bullet("build_mobilenetv2() - returns MobileNetV2 (ImageNet weights, include_top=False) -> GlobalAveragePooling2D -> Dense(256, ReLU) -> Dropout(0.5) -> Dense(6, Softmax). trainable_base defaults to False.")
add_bullet("build_resnet50() - same head as MobileNetV2, with ResNet50 backbone. trainable_base=False by default.")
add_bullet("build_rdf_random_forest() - returns a configured RandomForestClassifier (n_estimators=300, max_depth=10, min_samples_split=10, min_samples_leaf=4, class_weight='balanced_subsample', n_jobs=-1, random_state=42).")
add_para("Why these specific values? The RF hyperparameters come from the GridSearchCV run in train_rdf_rf.py (see best_params_ in results).")

add_heading("2.2.4 src/preprocessing.py - Image Preprocessing (615 lines)", level=3)
add_para("Five classes:")
add_bullet("ImageValidator - opens each image with PIL.Image.verify() and flags corrupted ones. Returns Dict[class_name] -> List[valid image paths]. For TrashNet: 2,527 valid, 0 corrupted.")
add_bullet("ImagePreprocessor - opens with PIL, converts to RGB, resizes to 224x224 with Lanczos interpolation, normalises pixel values to [0,1] (dividing by 255), stores as float32.")
add_bullet("DatasetSplitter - performs a stratified 70/15/15 split per class using sklearn's train_test_split. Two-step split: first separates 15% for test, then splits the remaining 85% into 70/15 train/val. Random state = 42.")
add_bullet("DatasetSaver - writes the preprocessed images as PNG to data/processed/{split}/{class}/{original_filename.png}.")
add_bullet("PreprocessingPipeline - orchestrator. Calls the four above in order, returns a results dictionary with valid_images_count, splits, saved_counts, class_distribution.")
add_para("Important nuance in DatasetSplitter: when splitting a single class, sklearn train_test_split is called WITHOUT a stratify argument (because there is only one class in that subset). This is per-class stratification rather than global stratification. The final result still preserves the global class distribution because the same train_ratio, val_ratio, test_ratio are applied to each class.")

add_heading("2.2.5 src/training.py - Training Utilities", level=3)
add_para("Defines TrainingManager class plus two helper functions and one RF trainer:")
add_bullet("TrainingManager.get_callbacks() - returns three Keras callbacks: EarlyStopping (monitor val_loss, patience 5, restore_best_weights=True), ModelCheckpoint (save best model to .h5, monitor val_accuracy, save_best_only=True), ReduceLROnPlateau (monitor val_loss, factor 0.5, patience 3, min_lr 1e-7).")
add_bullet("TrainingManager.compile_model() - supports 'adam', 'sgd' (with momentum 0.9), 'rmsprop'. Defaults to 'adam' with lr 0.001.")
add_bullet("TrainingManager.train() - calls model.fit() with the callbacks, returns the History object.")
add_bullet("create_data_loader() - wraps keras.preprocessing.image_dataset_from_directory. Returns a tf.data.Dataset. NOT used by the main scripts (they use image_dataset_from_directory directly), but available for reuse.")
add_bullet("prepare_dataset() - applies normalisation (divide by 255) and an optional preprocess_fn (e.g. mobilenet_preprocess_input or resnet_preprocess_input), then prefetches with AUTOTUNE. Used by compare_cnn_mobilenetv2.py.")
add_bullet("train_rdf_random_forest() - constructs a sklearn Pipeline(preprocessor, RandomForestClassifier) and runs GridSearchCV with 5-fold StratifiedKFold, scoring f1_weighted, refit=True. Default param grid: n_estimators [100, 200, 300], max_depth [10, 20, 30, None], min_samples_split [2, 5, 10], min_samples_leaf [1, 2, 4]. Returns the fitted GridSearchCV object.")

add_heading("2.2.6 src/evaluation.py - Metrics & Visualisation", level=3)
add_para("Four classes:")
add_bullet("MetricsCalculator - calculate_metrics(y_true, y_pred) returns {accuracy, precision, recall, f1_score} using sklearn with average='weighted'. get_confusion_matrix() and get_classification_report() wrap sklearn functions with the configured class_names.")
add_bullet("ConfusionMatrixVisualizer - static plot() method. Normalises by row (per-class percentage), draws a seaborn heatmap with annot showing percentages, axis labels 'Predicted Label' and 'True Label', 300 DPI.")
add_bullet("TrainingHistoryVisualizer - static plot_training_history() draws two figures: training+val accuracy and training+val loss, both with markers and grid. Saved to {filename_prefix}training_accuracy.png and ...loss.png.")
add_bullet("TabularModelVisualizer - static plot_feature_importance() draws a horizontal bar chart of importances from a trained Random Forest, sorted descending.")

add_heading("2.2.7 src/rdf_preprocessing.py - Tabular Pipeline", level=3)
add_para("RDFDataConfig - dataclass for CSV path, output dir, random_seed=42, n_samples=3000, test_size=0.2.")
add_para("RDFPreprocessingPipeline class:")
add_bullet("material_profiles - dictionary with six materials (cardboard, paper, plastic, metal, glass, organic), each with (min, max) ranges for moisture, contamination, combustibility, calorific.")
add_bullet("generate_dataset(n_samples=3000) - uses numpy's default_rng(seed=42) to sample uniformly within each material's range. Computes a synthetic rdf_score = 0.35*(calorific/46) + 0.25*(1-moisture/80) + 0.20*(combustibility/10) + 0.20*(1-contamination/10) + Gaussian noise. rdf_suitable = 1 if rdf_score >= 0.45 else 0. rdf_grade = High/Medium/Low/Unsuitable based on thresholds 0.7/0.5/0.35.")
add_bullet("save_dataset(df) - writes CSV to data/rdf_features/rdf_dataset.csv.")
add_bullet("load_or_generate_dataset() - loads if exists, else generates and saves.")
add_bullet("split_dataset(df) - sklearn train_test_split with test_size=0.2, stratify on rdf_suitable, random_state=42.")
add_bullet("build_preprocessor() - returns a ColumnTransformer: categorical pipeline (SimpleImputer most_frequent + OneHotEncoder handle_unknown=ignore, sparse_output=False) applied to material_type; numeric pipeline (SimpleImputer median + StandardScaler) applied to the other four columns.")

add_heading("2.2.8 src/multimodal_inference.py - The Chained Pipeline", level=3)
add_para("This is the heart of the multi-modal system. It chains an image classifier to a tabular RF predictor.")
add_bullet("ImageModelConfig - frozen dataclass with model_path, class_names (default from DEFAULT_DATA_CONFIG), input_shape (224,224), preprocess_mode ('baseline' | 'mobilenetv2' | 'resnet50').")
add_bullet("RDFModelConfig - just model_path.")
add_bullet("InferenceResult - frozen dataclass with image_path, predicted_class, class_confidence, material_features (dict), rdf_suitability (0/1), rdf_probability, rdf_label ('Suitable' or 'Not Suitable').")
add_bullet("MATERIAL_FEATURE_LIBRARY - hard-coded dictionary. Each waste class maps to: material_type, moisture_content, contamination_level, combustibility, calorific_value. Glass and metal have 0 combustibility and 0 calorific (they are not combustible). Trash maps to organic (60% moisture, 6 contamination, 4 combustibility, 6 calorific).")
add_bullet("ImageClassifier class - load_model (with custom_objects={'BaselineCNN': keras.Sequential} for compatibility), predict() loads image, resizes to 224, applies preprocess_fn (mobilenet_preprocess_input or resnet_preprocess_input or /255.0), calls model.predict(), returns (predicted_class, confidence, probabilities).")
add_bullet("MaterialFeatureMapper - build_features(waste_class) returns a single-row DataFrame with the 5 columns. Falls back to 'trash' if the class is not in the library.")
add_bullet("RDFSuitabilityPredictor - joblib.load, predict(features) returns (suitability, probability). If the model has predict_proba, the probability is class-1 probability; else defaults to 1.0.")
add_bullet("MultimodalInferencePipeline.infer(image_path) - the one-call function: image -> class -> material features -> RDF prediction -> InferenceResult.")

add_heading("2.2.9 src/utils.py - Stub", level=3)
add_para("Currently contains only a docstring listing intended utilities (logging, paths, configuration loading, device management, seed setting, file I/O). NO implementation. This is flagged as missing in the project analysis. In the viva, acknowledge that utils.py was a placeholder that the team did not have time to fill in.")

add_heading("2.3 The scripts/ Folder (Runners)", level=2)
add_para("These are the entry points you actually run from the command line. Each one wires up the src/ library and writes reports and figures.")
add_bullet("scripts/explore_dataset.py - Detects TrashNet/TACO, counts images, identifies corrupted files, generates reports/figures/trashnet_class_distribution.png and reports/dataset_report.md. Output: ~7 KB log.")
add_bullet("scripts/preprocess_dataset.py - Runs the PreprocessingPipeline. Produces data/processed/{train,validation,test}/{class}/, reports/figures/split_distribution.png, reports/figures/class_distribution_comparison.png, reports/preprocessing_report.md. Took ~18 seconds for 2,527 images.")
add_bullet("scripts/train_cnn.py - Trains the baseline CNN. Saves models/cnn_baseline_best.h5 (309 MB). Generates results/baseline_training_accuracy.png, baseline_training_loss.png, baseline_cnn_confusion_matrix.png, baseline_cnn_classification_report.txt, reports/cnn_baseline_report.md.")
add_bullet("scripts/compare_cnn_mobilenetv2.py - Trains baseline + MobileNetV2 + ResNet50 in one run. Uses two-phase training for the transfer-learning models (frozen backbone, then fine-tune top 30 layers at lr 1e-5). Saves three .h5 files, three confusion matrices, three classification reports, the comparison bar chart, and reports/cnn_mobilenetv2_resnet50_evaluation_report.md.")
add_bullet("scripts/train_rdf_rf.py - Generates or loads rdf_dataset.csv, runs GridSearchCV, saves models/rdf_random_forest_pipeline.joblib, generates results/rdf_confusion_matrix.png, results/rdf_feature_importance.png, results/rdf_classification_report.txt, reports/rdf_random_forest_report.md.")
add_bullet("scripts/run_multimodal_inference.py - CLI tool. Usage: python scripts/run_multimodal_inference.py --image PATH. Returns JSON with image_path, predicted_class, class_confidence, material_features, rdf_suitability, rdf_probability, rdf_label.")

add_heading("2.4 The data/ Folder", level=2)
add_bullet("data/processed/ - Gitignored. Contains 2,527 PNG images at 224x224 organised as train (1,763) / validation (381) / test (383) / {class}/.")
add_bullet("data/rdf_features/rdf_dataset.csv - 3,000 rows, 8 columns: material_type, moisture_content, contamination_level, combustibility, calorific_value, rdf_score, rdf_suitable (target), rdf_grade (ordinal).")
add_bullet("data/raw/ - Empty (raw is kept at TrashNET Data set/ in the repo root, not here).")

add_heading("2.5 The models/ Folder", level=2)
add_bullet("cnn_baseline_best.h5 - 309 MB. Keras HDF5 file. The trained baseline CNN.")
add_bullet("mobilenetv2_best.h5 - 25 MB. Smaller because MobileNetV2 is already a small architecture.")
add_bullet("resnet50_best.h5 - 216 MB. ResNet50 is larger than MobileNetV2 but smaller than baseline because the baseline has a huge Dense(256) layer with 25.6M parameters.")
add_bullet("rdf_random_forest_pipeline.joblib - 2.8 MB. scikit-learn Pipeline including the ColumnTransformer and the trained RandomForestClassifier.")

add_heading("2.6 The reports/ Folder (5 Markdown Files)", level=2)
add_bullet("dataset_report.md - Generated 2026-06-26. Documents TrashNet class distribution (cardboard 403, glass 501, metal 410, paper 594, plastic 482, trash 137 = 2,527 total). 0 corrupted. Resolution 512x384, all JPEG.")
add_bullet("preprocessing_report.md - Generated 2026-06-26. Documents 70/15/15 split, 1,763/381/383 totals, 224x224 RGB, [0,1] normalisation, 4.37:1 imbalance ratio.")
add_bullet("cnn_baseline_report.md - Generated 2026-07-02 18:54. The single-run baseline. 27 epochs, test acc 0.5614, F1 0.5505. NOTE: This is the FIRST baseline run; the COMPARISON run produced 0.5744 acc, 0.5723 F1 which supersedes it.")
add_bullet("cnn_mobilenetv2_resnet50_evaluation_report.md - Generated 2026-07-02 20:53. The canonical comparison. Baseline 0.5744, MobileNetV2 0.8198, ResNet50 0.8877.")
add_bullet("rdf_random_forest_report.md - Generated 2026-07-02 19:54. Best params, CV 0.9088, test acc 0.9133, F1 0.9141.")

add_heading("2.7 The results/ Folder (15 PNGs + 4 TXTs)", level=2)
add_bullet("baseline_cnn_classification_report.txt - per-class precision/recall/F1 for baseline CNN.")
add_bullet("mobilenetv2_classification_report.txt - per-class for MobileNetV2.")
add_bullet("resnet50_classification_report.txt - per-class for ResNet50 (the best image model).")
add_bullet("rdf_classification_report.txt - per-class for RDF (Not Suitable F1 0.9026, Suitable F1 0.9219).")
add_bullet("baseline_cnn_confusion_matrix.png, mobilenetv2_confusion_matrix.png, resnet50_confusion_matrix.png, rdf_confusion_matrix.png - 300 DPI heatmaps with percentage annotations.")
add_bullet("baseline_training_accuracy.png + baseline_training_loss.png - accuracy and loss curves for baseline CNN.")
add_bullet("mobilenetv2_training_accuracy.png + mobilenetv2_training_loss.png - same for MobileNetV2 (two phases concatenated).")
add_bullet("resnet50_training_accuracy.png + resnet50_training_loss.png - same for ResNet50.")
add_bullet("cnn_mobilenetv2_resnet50_comparison.png - grouped bar chart with 4 metrics x 3 models = 12 bars.")
add_bullet("rdf_feature_importance.png - bar chart of feature importances from the trained RF.")
add_bullet("Legacy duplicates: classification_report.txt, confusion_matrix.png, training_accuracy.png, training_loss.png. These predate the baseline_ prefix and should be cleaned up. (See recommendations.)")

add_heading("2.8 The docs/ Folder (7 Markdown Files)", level=2)
add_bullet("README.md - Index pointing to README.md, INSTALL.md, USAGE.md. (The latter two don't exist - be honest.)")
add_bullet("EXPLORATION_SCRIPT.md - Walkthrough of scripts/explore_dataset.py (230 lines).")
add_bullet("PREPROCESSING_PIPELINE.md - Detailed architecture of src/preprocessing.py (615 lines, 5 classes).")
add_bullet("PREPROCESSING_SUMMARY.md - Module-by-module explanation.")
add_bullet("PREPROCESSING_QUICK_REFERENCE.md - Cheatsheet.")
add_bullet("FLASK_WEB_APP_ARCHITECTURE.md - Design notes for app.py.")
add_bullet("MULTIMODAL_INFERENCE_ARCHITECTURE.md - Data flow for the chained pipeline.")

add_heading("2.9 templates/ and static/ (Flask UI)", level=2)
add_bullet("templates/index.html - Single-page Jinja2 template. Hero text, upload form, results section.")
add_bullet("static/style.css - Academic theme. CSS variables for colours (--accent: #12355b). Georgia serif font. Responsive grid with cards.")

add_heading("2.10 The TrashNET Data set/ Folder (Raw, Gitignored)", level=2)
add_para("The original TrashNet dataset. Inside dataset-resized/ you find six subfolders, one per class, with the raw 512x384 JPEG images. Total 2,527 images.")

add_heading("2.11 The TACO-master/ Folder", level=2)
add_para("TACO (Trash Annotations in Context) library cloned as a submodule. Contains annotations, download.py, detector/, demo.ipynb. The actual TACO images are NOT downloaded and NOT used. This is documented in the portfolio with an amber 'not used' badge.")

add_heading("2.12 The .venv/ Folder (Python Virtual Environment)", level=2)
add_para("Local virtual environment. Not tracked. Used to run the project locally on Windows.")

add_heading("2.13 Git History (5 Commits)", level=2)
add_bullet("565928c 2026-06-26 16:25 - Initialize MSc project structure (README, pyproject, .gitignore, requirements).")
add_bullet("91889f7 2026-06-26 16:26 - Implement dataset exploration and validation (scripts/explore_dataset.py).")
add_bullet("a42a9ff 2026-06-26 16:26 - Implement data preprocessing pipeline (src/preprocessing.py, scripts/preprocess_dataset.py).")
add_bullet("e631f51 2026-07-01 14:11 - Implement baseline CNN classifier (src/models.py, src/training.py, src/evaluation.py, scripts/train_cnn.py).")
add_bullet("500132a 2026-07-01 14:27 - Add multimodal inference and Flask app (src/multimodal_inference.py, src/rdf_preprocessing.py, scripts/*, app.py, templates, static, docs).")
add_para("Working tree state (at the time of writing): two uncommitted edits to src/multimodal_inference.py (custom_objects fix) and src/rdf_preprocessing.py (sparse -> sparse_output fix). The second was a real bug: the rdf_training.log shows the OneHotEncoder call failed twice on 2026-07-02 19:50 and 19:52 with 'unexpected keyword argument sparse' before being fixed at 19:54.")

add_heading("2.14 Data Flow Summary", level=2)
add_para("1. TrashNet raw images (TrashNET Data set/dataset-resized/) -> 2. explore_dataset.py validates and produces reports/dataset_report.md. -> 3. preprocess_dataset.py reads raw, preprocesses, splits 70/15/15, saves PNGs to data/processed/ and writes reports/preprocessing_report.md. -> 4. compare_cnn_mobilenetv2.py loads data/processed/, trains three models, evaluates on test set, writes reports and figures. -> 5. train_rdf_rf.py generates synthetic tabular data, trains RF, saves pipeline and report. -> 6. run_multimodal_inference.py takes any image, runs MultimodalInferencePipeline.infer(), prints JSON. -> 7. app.py wraps the same pipeline behind a Flask upload form.")

page_break()

# =====================================================================
# CHAPTER 3: SCRIPTS
# =====================================================================
add_heading("Chapter 3: Scripts (How to Run the Project)", level=1, color=RGBColor(0x12, 0x35, 0x5b))

add_heading("3.1 Why Scripts Exist", level=2)
add_para("Scripts are thin entry points. They wire together the src/ library, set up logging, manage directories, and generate reports. This separation means the heavy lifting (model definitions, preprocessing logic, evaluation) is testable and reusable, while the scripts are disposable one-off runners.")

add_heading("3.2 scripts/explore_dataset.py", level=3)
add_para("Purpose: walk a dataset, count images per class, check for corruption, log resolutions, generate a chart and a markdown report.")
add_para("Execution flow:")
add_numbered("setup_logging(exploration.log) -> console INFO + file DEBUG.")
add_numbered("DatasetDetector.detect_datasets() -> searches project root for 'TrashNET Data set' or 'TACO-master'.")
add_numbered("DatasetExplorer.explore() -> _collect_images_by_class() builds a defaultdict[class -> list[paths]]; _calculate_statistics() counts valid images per class, records resolution min/max/avg, lists corrupted paths.")
add_numbered("VisualizationGenerator -> saves 300 DPI PNG bar charts to reports/figures/.")
add_numbered("ReportGenerator -> writes reports/dataset_report.md.")
add_para("Input: file system only (TrashNET Data set/, TACO-master/).")
add_para("Output: exploration.log (7 KB), reports/dataset_report.md, reports/figures/trashnet_class_distribution.png, reports/figures/dataset_comparison.png.")
add_para("Generated files summary:")
add_bullet("reports/dataset_report.md (1,627 bytes).")
add_bullet("reports/figures/trashnet_class_distribution.png (117 KB).")
add_bullet("reports/figures/dataset_comparison.png (158 KB, same as class_distribution_comparison.png).")
add_bullet("exploration.log (~7 KB).")

add_heading("3.3 scripts/preprocess_dataset.py", level=3)
add_para("Purpose: take the raw images and produce a clean train/val/test directory ready for training.")
add_para("Execution flow:")
add_numbered("setup_logging(preprocessing.log).")
add_numbered("Construct PreprocessingPipeline(raw, processed, target_size=(224,224), ratios 0.7/0.15/0.15, seed=42).")
add_numbered("pipeline.run() -> validates -> splits -> saves PNGs.")
add_numbered("PreprocessingVisualizer -> saves split_distribution.png and class_distribution_comparison.png.")
add_numbered("PreprocessingReportGenerator -> writes reports/preprocessing_report.md.")
add_para("Input: TrashNET Data set/dataset-resized/{class}/{img}.jpg.")
add_para("Output: data/processed/{train,validation,test}/{class}/{img}.png (PNG, lossless, 224x224 RGB), reports/preprocessing_report.md, reports/figures/split_distribution.png, reports/figures/class_distribution_comparison.png, preprocessing.log.")
add_para("Key numbers (from preprocessing.log): cardboard 281/61/61, glass 350/75/76, metal 286/62/62, paper 415/89/90, plastic 336/73/73, trash 95/21/21. Total train 1,763 / val 381 / test 383.")

add_heading("3.4 scripts/train_cnn.py", level=3)
add_para("Purpose: train the baseline CNN end-to-end.")
add_para("Execution flow:")
add_numbered("setup_logging(training.log).")
add_numbered("load_preprocessed_datasets() -> uses keras.preprocessing.image_dataset_from_directory for train/val/test with image_size=(224,224), batch_size=32, label_mode='categorical', shuffle=True (train) / False (val, test), seed=42.")
add_numbered("prepare_datasets() -> maps a normalize function (images/255) and prefetches with AUTOTUNE.")
add_numbered("build_baseline_cnn() from src/models.py.")
add_numbered("TrainingManager.compile_model() with optimizer=adam, lr=0.001, loss=categorical_crossentropy, metrics=[accuracy].")
add_numbered("trainer.train() with epochs=30. Callbacks: EarlyStopping(patience=5, restore_best_weights), ModelCheckpoint(save_best_only), ReduceLROnPlateau(factor=0.5, patience=3).")
add_numbered("TrainingHistoryVisualizer.plot_training_history() -> baseline_training_accuracy.png and baseline_training_loss.png.")
add_numbered("evaluate_on_test_set() -> predict on test set, compute weighted accuracy/precision/recall/F1, save confusion matrix and classification report.")
add_numbered("generate_training_report() -> reports/cnn_baseline_report.md (this is the FIRST run, 0.5614 acc; the comparison run is the canonical one).")
add_para("Input: data/processed/.")
add_para("Output: models/cnn_baseline_best.h5 (309 MB), results/baseline_training_accuracy.png, results/baseline_training_loss.png, results/confusion_matrix.png (legacy duplicate), results/classification_report.txt (legacy duplicate), results/baseline_cnn_confusion_matrix.png, results/baseline_cnn_classification_report.txt, reports/cnn_baseline_report.md, training.log.")
add_para("Duration: training.log shows start 18:16:22, two re-runs (likely due to OOM), final completion 18:54:31 = ~38 minutes wall-clock for the successful 27-epoch run.")

add_heading("3.5 scripts/compare_cnn_mobilenetv2.py (3 models in one)", level=3)
add_para("Purpose: train and compare all three image models on the same split, then produce a side-by-side report.")
add_para("Execution flow:")
add_numbered("setup_logging(comparison.log).")
add_numbered("load_preprocessed_datasets() -> three tf.data.Datasets and the inferred class_names list.")
add_numbered("prepare_for_baseline() -> divide by 255.")
add_numbered("prepare_for_mobilenet() -> use mobilenet_preprocess_input, no manual normalisation.")
add_numbered("prepare_for_resnet() -> use resnet_preprocess_input, no manual normalisation.")
add_numbered("train_baseline() -> single phase, 30 epochs, callbacks as in train_cnn.py.")
add_numbered("train_mobilenetv2() -> PHASE 1: backbone frozen, 30 epochs, default lr 0.001. PHASE 2: unfreeze top 30 layers, recompile with lr=1e-5, train 10 more epochs. The two histories are concatenated.")
add_numbered("train_resnet50() -> same two-phase recipe.")
add_numbered("evaluate_model() for each -> metrics, confusion matrix saved to results/{name}_confusion_matrix.png, classification report saved to results/{name}_classification_report.txt.")
add_numbered("plot_comparison() -> 4-metric x 3-model grouped bar chart saved to results/cnn_mobilenetv2_resnet50_comparison.png.")
add_numbered("generate_comparison_report() -> reports/cnn_mobilenetv2_resnet50_evaluation_report.md.")
add_para("Final metrics (from comparison.log and reports):")
add_bullet("Baseline CNN: accuracy 0.5744, F1 0.5723.")
add_bullet("MobileNetV2: accuracy 0.8198, F1 0.8206.")
add_bullet("ResNet50: accuracy 0.8877, F1 0.8878.")
add_para("Input: data/processed/.")
add_para("Output: models/cnn_baseline_best.h5, models/mobilenetv2_best.h5, models/resnet50_best.h5, results/baseline_cnn_*.png/txt, results/mobilenetv2_*.png/txt, results/resnet50_*.png/txt, results/cnn_mobilenetv2_resnet50_comparison.png, reports/cnn_mobilenetv2_resnet50_evaluation_report.md, comparison.log.")
add_para("Duration: comparison.log shows the canonical run started 19:59:26 and completed 20:53:10 = ~54 minutes for all three models.")

add_heading("3.6 scripts/train_rdf_rf.py", level=3)
add_para("Purpose: build the RDF Random Forest pipeline and report.")
add_para("Execution flow:")
add_numbered("setup_logging(rdf_training.log).")
add_numbered("RDFPreprocessingPipeline(RDFDataConfig(csv_path=...)).")
add_numbered("load_or_generate_dataset() -> 3,000 synthetic rows.")
add_numbered("split_dataset(df) -> 2,400 train / 600 test, stratified on rdf_suitable.")
add_numbered("build_preprocessor() -> ColumnTransformer.")
add_numbered("train_rdf_random_forest() -> GridSearchCV over 3x4x3x3 = 108 combinations, 5-fold CV, f1_weighted scoring, n_jobs=-1.")
add_numbered("joblib.dump(search.best_estimator_, 'models/rdf_random_forest_pipeline.joblib').")
add_numbered("evaluate_model() -> predict on test, save rdf_confusion_matrix.png and rdf_classification_report.txt.")
add_numbered("get_feature_names_out() and classifier.feature_importances_ -> rdf_feature_importance.png.")
add_numbered("generate_report() -> reports/rdf_random_forest_report.md.")
add_para("Best params (from rdf_random_forest_report.md): n_estimators=300, max_depth=10, min_samples_split=10, min_samples_leaf=4. CV F1 (weighted) = 0.9088. Test acc 0.9133, F1 0.9141.")
add_para("Input: data/rdf_features/rdf_dataset.csv (or generated).")
add_para("Output: data/rdf_features/rdf_dataset.csv, models/rdf_random_forest_pipeline.joblib, results/rdf_confusion_matrix.png, results/rdf_classification_report.txt, results/rdf_feature_importance.png, reports/rdf_random_forest_report.md, rdf_training.log.")
add_para("Note: The rdf_training.log records that the first two runs (19:50 and 19:52) failed because OneHotEncoder no longer accepts 'sparse=False' in newer scikit-learn. The fix (sparse_output=False) made the third run (19:54) succeed.")

add_heading("3.7 scripts/run_multimodal_inference.py", level=3)
add_para("Purpose: a CLI tool for running the full chain on a single image.")
add_para("Usage:")
add_code("python scripts/run_multimodal_inference.py --image <path> [--image-model <.h5>] [--rdf-model <.joblib>] [--preprocess-mode baseline|mobilenetv2|resnet50] [--output <json>]")
add_para("Returns JSON to stdout with image_path, predicted_class, class_confidence, material_features, rdf_suitability (0/1), rdf_probability, rdf_label.")

add_heading("3.8 app.py (Flask)", level=3)
add_para("113 lines. Single route '/' with GET/POST. Uses functools.lru_cache(maxsize=1) to load the pipeline once. Validates file extension (.jpg/.jpeg/.png/.bmp/.gif/.tiff) and MIME prefix (image/). Max upload 16 MB. On POST: saves to NamedTemporaryFile, calls pipeline.infer(), encodes preview as base64, renders templates/index.html with the result.")
add_para("Run with: python app.py -> http://127.0.0.1:5000")

page_break()

# Save now to verify the build works so far
doc.save(OUT_PATH)
print(f"Part 1 saved to: {OUT_PATH}")
print(f"Current size: {os.path.getsize(OUT_PATH)} bytes")
